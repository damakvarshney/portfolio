import os
import subprocess

# Define paths
cwd = r"c:\Users\damak\Downloads\CodeNow\portfolio"
output_dir = os.path.join(cwd, "output", "pdf")
os.makedirs(output_dir, exist_ok=True)

brief_html_path = os.path.join(cwd, "output", "executive_brief.html")
ats_html_path = os.path.join(cwd, "output", "ats_resume.html")

brief_pdf_path = os.path.join(output_dir, "damak-varshney-executive-brief.pdf")
ats_pdf_path = os.path.join(output_dir, "damak-kumar-varshney-resume.pdf")

chrome_path = r"C:\Program Files\Google\Chrome\Application\chrome.exe"

# ----------------- 1. Executive Brief HTML & CSS -----------------
# This visual brief matches the website: deep navy black header, electric blue accents, visual timeline.
brief_html_content = """<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <title>Damak Kumar Varshney - Executive Brief</title>
  <link rel="preconnect" href="https://fonts.googleapis.com">
  <link rel="preconnect" href="https://fonts.gstatic.com" crossorigin>
  <link href="https://fonts.googleapis.com/css2?family=DM+Sans:wght@300;400;500;700&family=PT+Sans:wght@400;700&display=swap" rel="stylesheet">
  <style>
    @page {
      size: letter;
      margin: 0;
    }
    *, *::before, *::after {
      box-sizing: border-box;
    }
    body {
      font-family: 'DM Sans', sans-serif;
      color: #0F172A;
      background-color: #FFFFFF;
      margin: 0;
      padding: 0;
      -webkit-print-color-adjust: exact;
      font-size: 8pt;
      line-height: 1.25;
    }
    .page {
      width: 8.5in;
      height: 11in;
      position: relative;
      box-sizing: border-box;
      background-color: #FFFFFF;
      overflow: hidden;
    }
    .p-padded {
      padding: 0.25in 0.45in;
    }
    .header-content {
      background-color: #080C10;
      color: #FFFFFF;
      padding: 0.25in 0.45in;
    }
    .header-content h1 {
      font-family: 'PT Sans', sans-serif;
      font-size: 20pt;
      margin: 0 0 2px 0;
      font-weight: 700;
      letter-spacing: -0.5px;
      color: #FFFFFF;
    }
    .header-content .subtitle {
      color: #2563EB;
      font-family: 'PT Sans', sans-serif;
      font-size: 10pt;
      font-weight: 700;
      margin: 0 0 6px 0;
      letter-spacing: 0.5px;
      text-transform: uppercase;
    }
    .header-content .contacts-grid {
      display: flex;
      flex-wrap: wrap;
      gap: 5px 15px;
      font-size: 7.5pt;
      color: #94A3B8;
    }
    .header-content .contacts-grid span {
      display: flex;
      align-items: center;
      gap: 3px;
    }
    .header-content .contacts-grid a {
      color: #94A3B8;
      text-decoration: none;
    }

    /* Sections */
    .section-title {
      font-family: 'PT Sans', sans-serif;
      font-size: 9.5pt;
      font-weight: 700;
      color: #2563EB;
      border-bottom: 1px solid #CBD5E1;
      padding-bottom: 2px;
      margin-top: 8px;
      margin-bottom: 4px;
      text-transform: uppercase;
      letter-spacing: 0.5px;
    }

    /* Positioning */
    .positioning-text {
      font-size: 8pt;
      color: #334155;
      line-height: 1.35;
      margin-bottom: 6px;
    }

    /* Metrics Strip */
    .metrics-grid {
      display: grid;
      grid-template-columns: repeat(6, 1fr);
      gap: 6px;
      margin-bottom: 6px;
    }
    .metric-box {
      background-color: #F8FAFC;
      border: 1px solid #CBD5E1;
      border-radius: 6px;
      padding: 4px 2px;
      text-align: center;
    }
    .metric-value {
      font-size: 12pt;
      font-weight: 700;
      color: #2563EB;
      font-family: 'PT Sans', sans-serif;
      line-height: 1.1;
    }
    .metric-label {
      font-size: 5.5pt;
      font-weight: 700;
      color: #0F172A;
      margin-top: 1px;
      text-transform: uppercase;
      letter-spacing: 0.1px;
      line-height: 1;
    }

    /* Operating Scale */
    .scale-grid {
      display: grid;
      grid-template-columns: repeat(3, 1fr);
      gap: 8px;
      background-color: #F8FAFC;
      border: 1px solid #E2E8F0;
      border-radius: 6px;
      padding: 6px 12px;
      margin-bottom: 6px;
    }
    .scale-col {
      font-size: 7.5pt;
      color: #334155;
      line-height: 1.3;
    }
    .scale-item {
      display: flex;
      justify-content: space-between;
    }
    .scale-item span:first-child {
      font-weight: 700;
      color: #0F172A;
    }

    /* Case Studies 2x3 Grid */
    .case-grid {
      display: grid;
      grid-template-columns: 1fr 1fr;
      gap: 6px;
      margin-bottom: 6px;
    }
    .case-card {
      border: 1px solid #E2E8F0;
      border-left: 2.5px solid #2563EB;
      border-radius: 5px;
      padding: 5px 8px;
      background-color: #FCFDFE;
    }
    .case-header {
      display: flex;
      justify-content: space-between;
      align-items: center;
      margin-bottom: 2px;
    }
    .case-title {
      font-family: 'PT Sans', sans-serif;
      font-size: 8pt;
      font-weight: 700;
      color: #0F172A;
    }
    .case-tag {
      font-size: 6pt;
      font-weight: 700;
      color: #2563EB;
      background-color: #EFF6FF;
      padding: 1px 4px;
      border-radius: 3px;
      border: 0.5px solid #DBEAFE;
      white-space: nowrap;
    }
    .case-detail {
      font-size: 7.2pt;
      color: #475569;
      margin-bottom: 1px;
      line-height: 1.2;
    }
    .case-impact {
      font-size: 7.2pt;
      font-weight: 700;
      color: #10B981;
      line-height: 1.1;
    }

    /* Capability Stack */
    .stack-container {
      display: flex;
      flex-direction: column;
      gap: 3px;
      margin-bottom: 6px;
    }
    .stack-row {
      display: grid;
      grid-template-columns: 1.4in 1fr;
      gap: 10px;
      font-size: 7.5pt;
      line-height: 1.2;
    }
    .stack-label {
      font-weight: 700;
      color: #0F172A;
    }
    .stack-value {
      color: #334155;
    }

    /* Recognition Strip */
    .recog-container {
      display: flex;
      flex-direction: column;
      gap: 2px;
      margin-bottom: 6px;
    }
    .recog-item {
      font-size: 7.5pt;
      color: #334155;
      line-height: 1.2;
    }
    .recog-bullet {
      color: #2563EB;
      font-weight: bold;
      margin-right: 3px;
    }

    /* Education Footer */
    .edu-footer {
      border-top: 1px solid #CBD5E1;
      padding-top: 4px;
      margin-top: 6px;
      display: flex;
      justify-content: space-between;
      font-size: 7.5pt;
      color: #475569;
    }
    .edu-item {
      font-weight: 500;
    }
    .edu-item strong {
      color: #0F172A;
    }
  </style>
</head>
<body>

  <div class="page">
    <!-- Header -->
    <div class="header-content">
      <h1>DAMAK KUMAR VARSHNEY</h1>
      <div class="subtitle">AI Product Manager | Business Transformation | Enterprise AI Strategy</div>
      <div class="contacts-grid">
        <span>📍 Gurugram, India</span>
        <span>📞 +91-8077293920</span>
        <span>✉️ <a href="mailto:damakvarshney@gmail.com">damakvarshney@gmail.com</a></span>
        <span>🔗 <a href="https://linkedin.com/in/damak-varshney" target="_blank">linkedin.com/in/damak-var-shney</a></span>
        <span>🌐 <a href="https://damak-varshney.vercel.app" target="_blank">damak-varshney.vercel.app</a></span>
      </div>
    </div>

    <div class="p-padded">
      <!-- Positioning -->
      <div class="positioning-text">
        <strong>POSITIONING:</strong> AI Product Manager with 2+ years driving enterprise AI transformation at OSTTRA (KKR Portfolio). Architect of AI systems processing 200,000+ monthly cases with &gt;99% accuracy. Designer of prompts powering sentiment analytics across 20,000+ client communications. Combines technical build depth, financial-services expertise, and enterprise adoption discipline.
      </div>

      <!-- Impact Metrics -->
      <div class="metrics-grid">
        <div class="metric-box">
          <div class="metric-value">20+</div>
          <div class="metric-label">FTE Saved</div>
        </div>
        <div class="metric-box">
          <div class="metric-value">200K+</div>
          <div class="metric-label">Cases Monthly</div>
        </div>
        <div class="metric-box">
          <div class="metric-value">1,500+</div>
          <div class="metric-label">Users</div>
        </div>
        <div class="metric-box">
          <div class="metric-value">55%</div>
          <div class="metric-label">Faster KYC</div>
        </div>
        <div class="metric-box">
          <div class="metric-value">&gt;99%</div>
          <div class="metric-label">Accuracy</div>
        </div>
        <div class="metric-box">
          <div class="metric-value">40%</div>
          <div class="metric-label">Faster Res.</div>
        </div>
      </div>

      <!-- Operating Scale -->
      <div class="section-title">Operating at Enterprise Scale</div>
      <div class="scale-grid">
        <div class="scale-col">
          <div class="scale-item"><span>200,000+</span> <span>cases routed monthly</span></div>
          <div class="scale-item"><span>4,500+</span> <span>system alerts handled</span></div>
        </div>
        <div class="scale-col" style="border-left: 1px solid #E2E8F0; border-right: 1px solid #E2E8F0; padding: 0 8px;">
          <div class="scale-item"><span>20,000+</span> <span>emails analyzed</span></div>
          <div class="scale-item"><span>751</span> <span>screenings monthly</span></div>
        </div>
        <div class="scale-col">
          <div class="scale-item"><span>71</span> <span>AI initiatives delivered</span></div>
          <div class="scale-item"><span>1,500+</span> <span>users at 80% adoption</span></div>
        </div>
      </div>

      <!-- Case Studies -->
      <div class="section-title">Selected Case Studies</div>
      <div class="case-grid">
        <!-- Card 1 -->
        <div class="case-card">
          <div class="case-header">
            <span class="case-title">NEXUS — Global Capacity Intelligence</span>
            <span class="case-tag">Personal Lead</span>
          </div>
          <div class="case-detail">Before: 150+ spreadsheets &rarr; After: Real-time AI platform</div>
          <div class="case-impact">Impact: 80%+ adoption | 1,500+ users | 45% faster planning</div>
        </div>
        <!-- Card 2 -->
        <div class="case-card">
          <div class="case-header">
            <span class="case-title">Auto-Triaging Ecosystem</span>
            <span class="case-tag">Personal Lead</span>
          </div>
          <div class="case-detail">Before: Manual L1 triage &rarr; After: Autonomous routing</div>
          <div class="case-impact">Impact: 200K+ cases | &gt;99% accuracy | 8.78 FTE reclaimed</div>
        </div>
        <!-- Card 3 -->
        <div class="case-card">
          <div class="case-header">
            <span class="case-title">KYC & Onboarding Automation</span>
            <span class="case-tag">Personal Lead</span>
          </div>
          <div class="case-detail">Before: 18-day manual vetting &rarr; After: 8-day automated cycle</div>
          <div class="case-impact">Impact: 55% faster | 751 screenings/mo | 1.0 FTE reclaimed</div>
        </div>
        <!-- Card 4 -->
        <div class="case-card">
          <div class="case-header">
            <span class="case-title">Autonomous Failure Management</span>
            <span class="case-tag">Personal Lead</span>
          </div>
          <div class="case-detail">Before: Manual alert triage &rarr; After: Self-routing system</div>
          <div class="case-impact">Impact: 4,500+ alerts/mo | Zero human triage required</div>
        </div>
        <!-- Card 5 -->
        <div class="case-card">
          <div class="case-header">
            <span class="case-title">Predictive Sentiment Analytics</span>
            <span class="case-tag">Prompt + Change Lead</span>
          </div>
          <div class="case-detail">Before: Reactive escalation &rarr; After: Proactive outreach</div>
          <div class="case-impact">Impact: 20K+ emails | 40% faster res. | 20% fewer escalations</div>
        </div>
        <!-- Card 6 -->
        <div class="case-card">
          <div class="case-header">
            <span class="case-title">Dot, ODA & Transformer Bot</span>
            <span class="case-tag">Contributing Member</span>
          </div>
          <div class="case-detail">Before: Manual queries &rarr; After: RAG-powered self-service</div>
          <div class="case-impact">Impact: 16% to 3% cases | 4+ FTE freed | 24 initiatives</div>
        </div>
      </div>

      <!-- Capability Stack -->
      <div class="section-title">Capability Stack</div>
      <div class="stack-container">
        <div class="stack-row">
          <div class="stack-label">▪ AI / LLM Stack</div>
          <div class="stack-value">Vertex AI · Gemini · RAG · Prompt Engineering · Conversational AI · Predictive Sentiment · Agentic Workflows</div>
        </div>
        <div class="stack-row">
          <div class="stack-label">▪ Cloud & Data</div>
          <div class="stack-value">Python · SQL · BigQuery · ChromaDB · GCP · Looker · Power BI · Data Analytics · Business Intelligence</div>
        </div>
        <div class="stack-row">
          <div class="stack-label">▪ Product Management</div>
          <div class="stack-value">Product Strategy · Roadmapping · User Adoption · KPI Tracking · Cross-Functional Leadership</div>
        </div>
        <div class="stack-row">
          <div class="stack-label">▪ Methodologies</div>
          <div class="stack-value">Lean Six Sigma · Agile/Scrum · LUMA Design Thinking · Change Management · Continuous Improvement</div>
        </div>
        <div class="stack-row">
          <div class="stack-label">▪ Financial Domain</div>
          <div class="stack-value">OTC Markets · Reference Data · KYC/AML · SOX · Trade Failure Management · Regulatory Reporting</div>
        </div>
      </div>

      <!-- Recognition -->
      <div class="section-title">Recognition</div>
      <div class="recog-container">
        <div class="recog-item"><span class="recog-bullet">★</span> Times of India Feature — InspireTech 2025 (Enterprise AI & Automation)</div>
        <div class="recog-item"><span class="recog-bullet">★</span> Highflyer Award — OSTTRA Q1 2026 (Looker & Auto-Triaging Impact)</div>
        <div class="recog-item"><span class="recog-bullet">★</span> Trendsetter Award — OSTTRA Q1 2025 (SOX Compliance Automation)</div>
        <div class="recog-item"><span class="recog-bullet">★</span> 1st Prize — LUMA Design Thinking Challenge, Protiviti 2023</div>
        <div class="recog-item"><span class="recog-bullet">★</span> Multiple Quarterly Spot Awards — OSTTRA 2024–2026</div>
      </div>

      <!-- Education Footer -->
      <div class="edu-footer">
        <div class="edu-item"><strong>PGPM Finance & Analytics</strong>, IBS Gurgaon (CGPA 9.0/10, 2024)</div>
        <div class="edu-item"><strong>MBA Finance & Analytics</strong>, IFHE Hyderabad (2024)</div>
        <div class="edu-item"><strong>BCA Computer Science</strong>, BIT Mesra (2020)</div>
      </div>
    </div>
  </div>

</body>
</html>
"""

# ----------------- 2. ATS-Optimized Resume HTML & CSS -----------------
# Fully revised: tighter margins, phone number, no arrows, varied verbs, semicolons, compact bullets, pipe-separated skills, correct section order.
ats_html_content = """<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="UTF-8">
  <title>Damak Kumar Varshney - Resume</title>
  <style>
    @page {
      size: letter;
      margin: 0.25in 0.45in;
      /* Suppress all browser headers and footers */
      @top-left { content: ""; }
      @top-center { content: ""; }
      @top-right { content: ""; }
      @bottom-left { content: ""; }
      @bottom-center { content: ""; }
      @bottom-right { content: ""; }
    }
    @media print {
      body { margin: 0; }
      .no-print { display: none; }
    }
    body {
      font-family: 'Calibri', 'Arial', sans-serif;
      color: #000000;
      background-color: #FFFFFF;
      margin: 0;
      padding: 0;
      font-size: 8.5pt;
      line-height: 1.05;
    }
    .header-section {
      text-align: center;
      margin-bottom: 3px;
    }
    .name {
      font-size: 13pt;
      font-weight: bold;
      text-transform: uppercase;
      margin-bottom: 1px;
      letter-spacing: 0.5px;
    }
    .contact-info {
      font-size: 8pt;
      margin-bottom: 1px;
    }
    .headline {
      font-size: 8.5pt;
      font-weight: bold;
      text-transform: uppercase;
      border-bottom: 1px solid #000000;
      padding-bottom: 1px;
      margin-top: 1px;
      margin-bottom: 3px;
    }
    .section-title {
      font-size: 9pt;
      font-weight: bold;
      text-transform: uppercase;
      border-bottom: 1px solid #000000;
      margin-top: 4px;
      margin-bottom: 2px;
      padding-bottom: 1px;
    }
    .summary-text {
      font-size: 8.2pt;
      margin-bottom: 0;
      text-align: justify;
      line-height: 1.1;
    }
    .job-entry {
      margin-bottom: 2px;
    }
    .job-header {
      font-weight: bold;
      font-size: 8.5pt;
      display: flex;
      justify-content: space-between;
    }
    .job-meta {
      font-style: italic;
      font-size: 7.8pt;
      display: flex;
      justify-content: space-between;
      margin-bottom: 1px;
    }
    ul {
      margin: 0px;
      padding-left: 10px;
    }
    li {
      margin-bottom: 0px;
      text-align: justify;
      font-size: 8.2pt;
      line-height: 1.05;
    }
    .education-line {
      font-size: 8.2pt;
      margin-bottom: 0px;
      line-height: 1.05;
    }
    .skills-category {
      margin-bottom: 0px;
      font-size: 8.2pt;
      line-height: 1.05;
    }
    .skills-label {
      font-weight: bold;
    }
    .cert-line {
      font-size: 8.2pt;
      margin-bottom: 0px;
      line-height: 1.05;
    }
    .recog-list {
      margin: 0px;
        <style>
    @page {
      size: letter;
      margin: 0.25in 0.45in;
      /* Suppress all browser headers and footers */
      @top-left { content: ""; }
      @top-center { content: ""; }
      @top-right { content: ""; }
      @bottom-left { content: ""; }
      @bottom-center { content: ""; }
      @bottom-right { content: ""; }
    }
    @media print {
      body { margin: 0; }
      .no-print { display: none; }
    }
    body {
      font-family: 'Calibri', 'Arial', sans-serif;
      color: #000000;
      background-color: #FFFFFF;
      margin: 0;
      padding: 0;
      font-size: 8.5pt;
      line-height: 1.05;
    }
    .header-section {
      text-align: center;
      margin-bottom: 3px;
    }
    .name {
      font-size: 13pt;
      font-weight: bold;
      text-transform: uppercase;
      margin-bottom: 1px;
      letter-spacing: 0.5px;
    }
    .contact-info {
      font-size: 8pt;
      margin-bottom: 1px;
    }
    .headline {
      font-size: 8.5pt;
      font-weight: bold;
      text-transform: uppercase;
      border-bottom: 1px solid #000000;
      padding-bottom: 1px;
      margin-top: 1px;
      margin-bottom: 3px;
    }
    .section-title {
      font-size: 9pt;
      font-weight: bold;
      text-transform: uppercase;
      border-bottom: 1px solid #000000;
      margin-top: 4px;
      margin-bottom: 2px;
      padding-bottom: 1px;
    }
    .summary-text {
      font-size: 8.2pt;
      margin-bottom: 0;
      text-align: justify;
      line-height: 1.1;
    }
    .job-entry {
      margin-bottom: 2px;
    }
    .job-header {
      font-weight: bold;
      font-size: 8.5pt;
      display: flex;
      justify-content: space-between;
    }
    .job-meta {
      font-style: italic;
      font-size: 7.8pt;
      display: flex;
      justify-content: space-between;
      margin-bottom: 1px;
    }
    ul {
      margin: 0px;
      padding-left: 0px;
      list-style-type: none;
    }
    li {
      margin-bottom: 0px;
      text-align: justify;
      font-size: 8.2pt;
      line-height: 1.05;
      padding-left: 8px;
      text-indent: -8px;
    }
    .education-line {
      font-size: 8.2pt;
      margin-bottom: 0px;
      line-height: 1.05;
    }
    .skills-category {
      margin-bottom: 0px;
      font-size: 8.2pt;
      line-height: 1.05;
    }
    .skills-label {
      font-weight: bold;
    }
    .cert-line {
      font-size: 8.2pt;
      margin-bottom: 0px;
      line-height: 1.05;
    }
    .recog-list {
      margin: 0px;
      padding-left: 10px;
      list-style-type: none;
    }
    .recog-list li {
      font-size: 8.2pt;
      margin-bottom: 0px;
      line-height: 1.05;
    }
  </style>
</head>
<body>

  <div class="header-section">
    <div class="name">DAMAK KUMAR VARSHNEY</div>
    <div class="contact-info">
      Gurugram, India &nbsp;|&nbsp; +91-8077293920 &nbsp;|&nbsp; damakvarshney@gmail.com
    </div>
    <div class="contact-info">
      linkedin.com/in/damak-varshney &nbsp;|&nbsp; damak-varshney.vercel.app
    </div>
    <div class="headline">
      AI PRODUCT MANAGER &nbsp;|&nbsp; BUSINESS TRANSFORMATION &nbsp;|&nbsp; ENTERPRISE AI STRATEGY
    </div>
  </div>

  <div class="section-title">SUMMARY</div>
  <div class="summary-text">
    AI Product Manager with 2+ years driving enterprise AI transformation at OSTTRA, a KKR portfolio company, and 4+ years total in enterprise automation and product delivery. Architect of AI systems processing 200,000+ monthly cases with &gt;99% routing accuracy and designed prompts powering predictive sentiment analytics across 20,000+ client communications. Delivered 20+ FTE in annualized savings and 80%+ adoption across 1,500+ enterprise users through Product Strategy, Change Management, and rigorous adoption discipline. Expertise spans Vertex AI, RAG pipelines, Prompt Engineering, financial-services automation, and data-driven decision making. Open to AI Product Manager, AI Strategy, and Enterprise Transformation roles.
  </div>

  <div class="section-title">WORK EXPERIENCE</div>

  <div class="job-entry">
    <div class="job-header">
      <span>Associate, AI Strategy &amp; Ops Excellence (DigiOps Programme)</span>
      <span>Apr 2025 &ndash; Present</span>
    </div>
    <div class="job-meta">
      <span>OSTTRA (KKR Portfolio Company) | Gurugram, India</span>
    </div>
    <ul>
      <li>• Architected NEXUS global capacity intelligence platform consolidating 150+ spreadsheets using Python predictive modeling, Business Intelligence dashboards via Looker, and Data Analytics pipelines; drove 80%+ adoption across 1,500+ users and compressed planning cycles 45%.</li>
      <li>• Spearheaded Generative AI-powered Auto-Triaging ecosystem processing 200,000+ monthly cases with &gt;99% routing accuracy; reclaimed 4.9 FTEs through Process Automation and scaled duplicate flagging to 15,206 monthly checks recovering 3.88 additional FTEs.</li>
      <li>• Led intelligent KYC and onboarding automation through API Integration with reference registry systems executing 751 monthly compliance screenings; shortened cycle time 55% from 18 to 8 days, reclaimed 1.0 FTE, and established 100% audit traceability.</li>
      <li>• Engineered autonomous failure management for 4,500+ monthly complex system alerts including LimitHub and TOC FIX disconnections; eliminated reactive triage burden by routing non-actionable items without human intervention.</li>
      <li>• Designed AI prompts and led change management for predictive sentiment analytics across 20,000+ client emails on MarkitWire and TriResolve; engineered dual-category classification (positive, negative, neutral) for Client Services and Product, enabling proactive client outreach before escalation accelerating resolution 40% and decreasing escalations 20%.</li>
      <li>• Contributed to the design and rollout of Dot, ODA, and Transformer Bot, RAG-based Conversational AI assistants powered by Vertex AI; supported the team that cut routine client information cases from 16% to 3%, freed 4+ FTE annually, and deployed across 24 knowledge initiatives.</li>
    </ul>
  </div>

  <div class="job-entry">
    <div class="job-header">
      <span>Senior Operational Excellence Analyst</span>
      <span>Feb 2024 &ndash; Dec 2024</span>
    </div>
    <div class="job-meta">
      <span>OSTTRA | Gurugram, India</span>
    </div>
    <ul>
      <li>• Engineered SOX Compliance automation using Google Apps Script; eliminated 4.1 FTEs of manual audit work, achieved 100% compliance, and compressed audit cycle from 45 to 5 days.</li>
      <li>• Delivered 10+ continuous improvement initiatives applying Lean Six Sigma and Agile across client services and trading operations; unlocked 2+ FTE in quantitative savings.</li>
      <li>• Directed Microsoft to Google Workspace migration during Tel Aviv-to-India transition; removed VDI dependency, cut access latency 40%, and freed 0.4 FTEs.</li>
    </ul>
  </div>

  <div class="job-entry">
    <div class="job-header">
      <span>Credit Analyst Intern</span>
      <span>Feb 2023 &ndash; Apr 2023</span>
    </div>
    <div class="job-meta">
      <span>Kotak Mahindra Bank | Delhi NCR, India</span>
    </div>
    <ul>
      <li>• Evaluated 20+ MSME credit applications analyzing financial statements and working capital dynamics; delivered recommendations with 95%+ accuracy contributing to $3M+ underwriting decisions.</li>
    </ul>
  </div>

  <div class="job-entry">
    <div class="job-header">
      <span>React Native Developer</span>
      <span>Apr 2021 &ndash; May 2022</span>
    </div>
    <div class="job-meta">
      <span>RPQ IT Services | Noida (Sector 16), India</span>
    </div>
    <ul>
      <li>• Developed 5+ production mobile applications using React Native with advanced state management and component architecture.</li>
      <li>• Managed full Software Development Life Cycle (SDLC) from requirement gathering through production deployment for iOS and Android platforms.</li>
      <li>• Architected cross-platform build pipelines and CI/CD integration enabling multi-platform releases.</li>
    </ul>
  </div>

  <div class="job-entry">
    <div class="job-header">
      <span>React Native Developer</span>
      <span>Nov 2020 &ndash; Mar 2021</span>
    </div>
    <div class="job-meta">
      <span>Markup Designer | Noida (Sector 62), India</span>
    </div>
    <ul>
      <li>• Engineered React Native mobile applications with REST API integration and asynchronous data handling.</li>
      <li>• Collaborated with design teams using Figma and Adobe XD translating wireframes into production-grade cross-platform code.</li>
      <li>• Optimized application performance achieving 40%+ faster load times through code splitting and lazy loading.</li>
    </ul>
  </div>

  <div class="job-entry">
    <div class="job-header">
      <span>Mobile Application Developer</span>
      <span>Jul 2020 &ndash; Nov 2020</span>
    </div>
    <div class="job-meta">
      <span>Sabpay | Remote, India (BCA Freelance Engagement)</span>
    </div>
    <ul>
      <li>• Built mobile application user interfaces using React Native during freelance engagement while completing BCA.</li>
      <li>• Delivered UI/UX prototypes in Figma and Adobe XD shipping production builds on schedule.</li>
      <li>• Implemented OAuth authentication and payment gateway integration within the Sabpay ecosystem.</li>
    </ul>
  </div>

  <div class="section-title">EDUCATION</div>
  <div class="education-line"><strong>Post Graduate Programme, Finance &amp; Analytics</strong> &mdash; ICFAI Business School, Gurgaon | 2022 &ndash; Apr 2024 | CGPA: 9.0/10</div>
  <div class="education-line"><strong>MBA, Finance &amp; Analytics</strong> (Online, concurrent with OSTTRA role) &mdash; IFHE Hyderabad | 2022 &ndash; Jul 2024 | CGPA: 8.75/10</div>
  <div class="education-line"><strong>Bachelor of Computer Applications</strong> &mdash; Birla Institute of Technology, Mesra | 2017 &ndash; 2020 | CGPA: 7.52/10</div>

  <div class="section-title">SKILLS</div>
  <div class="skills-category"><span class="skills-label">AI and LLM Stack:</span> Vertex AI, Google Gemini, Generative AI (GenAI), Machine Learning, RAG Pipelines, Vector Embeddings, Prompt Engineering, LLM Optimization, Conversational AI, Predictive Sentiment Analysis, Agentic Workflows, LLM-Powered Automation, ChatGPT</div>
  <div class="skills-category"><span class="skills-label">Cloud and Data:</span> Python, SQL, Google BigQuery, ChromaDB, Google Cloud Platform (GCP), Looker Studio, Power BI, Data Analytics, Business Intelligence, Data Engineering, Metrics-Driven Analysis</div>
  <div class="skills-category"><span class="skills-label">Product and Tools:</span> Salesforce (Administration, Dashboards, Case Management, Custom Fields), Jira, Confluence, Google AppSheet, Google Apps Script, REST APIs, API Integration</div>
  <div class="skills-category"><span class="skills-label">Product Management:</span> Product Strategy, Requirements Definition, Roadmap Development, User Adoption, Feature Prioritization, Sprint Planning, KPI Tracking, Stakeholder Communication, Executive Reporting, Cross-Functional Leadership, End-to-End Ownership</div>
  <div class="skills-category"><span class="skills-label">Methodologies:</span> Lean Six Sigma, Agile/Scrum, LUMA Design Thinking, Process Reengineering, Change Management, Stakeholder Enablement, Continuous Improvement, Data-Driven Decision Making, Cost Optimization</div>
  <div class="skills-category"><span class="skills-label">Financial Domain:</span> OTC Markets, Reference Data, Trade Operations, Settlement, Reconciliation, KYC/AML, SOX Compliance, Regulatory Reporting, Audit Trail Management, Trade Failure Management, Financial Services Infrastructure, Credit Analysis</div>

  <div class="section-title">CERTIFICATIONS</div>
  <div class="cert-line">Google Data Analytics Professional Certificate, Google (2022)</div>
  <div class="cert-line">Certified Data Scientist, Henry Harvin Education (2022)</div>
  <div class="cert-line">SQL Certification, DataCamp (2022)</div>

  <div class="section-title">RECOGNITION</div>
  <ul class="recog-list">
    <li>• Featured in Times of India, InspireTech 2025 &mdash; Enterprise AI and Automation Systems</li>
    <li>• Highflyer Award, OSTTRA Q1 2026 &mdash; Looker analytics and Salesforce auto-triaging impact</li>
    <li>• Trendsetter Award, OSTTRA Q1 2025 &mdash; SOX Compliance automation</li>
    <li>• 1st Prize, LUMA Design Thinking Challenge, Protiviti 2023</li>
    <li>• Multiple Quarterly Spot Awards, OSTTRA 2024 &ndash; 2026</li>
  </ul>

</body>
</html>
"""

# Write HTML files
with open(brief_html_path, "w", encoding="utf-8") as f:
    f.write(brief_html_content)
print(f"Written {brief_html_path}")

with open(ats_html_path, "w", encoding="utf-8") as f:
    f.write(ats_html_content)
print(f"Written {ats_html_path}")

# Compile HTML to PDF using headless Chrome
print("Compiling Executive Brief to PDF...")
brief_cmd = [
    chrome_path,
    "--headless",
    "--disable-gpu",
    "--no-header-footer",
    f"--print-to-pdf={brief_pdf_path}",
    brief_html_path
]
subprocess.run(brief_cmd, check=True)
print(f"Successfully generated {brief_pdf_path}")

print("Compiling ATS Resume to PDF...")
ats_cmd = [
    chrome_path,
    "--headless",
    "--disable-gpu",
    "--no-header-footer",
    f"--print-to-pdf={ats_pdf_path}",
    ats_html_path
]
subprocess.run(ats_cmd, check=True)
print(f"Successfully generated {ats_pdf_path}")

# Generate DOCX Version
def build_docx(output_path):
    import docx
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    
    # 0.75 margin all sides
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Base style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    
    def format_spacing(p, space_before=0, space_after=2, line_spacing=1.1):
        p_format = p.paragraph_format
        p_format.space_before = Pt(space_before)
        p_format.space_after = Pt(space_after)
        p_format.line_spacing = line_spacing

    def add_bottom_border(paragraph):
        pPr = paragraph._p.get_or_add_pPr()
        pbdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pbdr.append(bottom)
        pPr.append(pbdr)

    # Header
    p_name = doc.add_paragraph()
    format_spacing(p_name, 0, 1)
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run("DAMAK KUMAR VARSHNEY")
    run_name.font.size = Pt(15)
    run_name.bold = True

    p_contact1 = doc.add_paragraph()
    format_spacing(p_contact1, 0, 1)
    p_contact1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_c1 = p_contact1.add_run("Gurugram, India  |  +91-8077293920  |  damakvarshney@gmail.com")
    run_c1.font.size = Pt(9.5)

    p_contact2 = doc.add_paragraph()
    format_spacing(p_contact2, 0, 2)
    p_contact2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_c2 = p_contact2.add_run("linkedin.com/in/damak-varshney  |  damak-varshney.vercel.app")
    run_c2.font.size = Pt(9.5)

    p_headline = doc.add_paragraph()
    format_spacing(p_headline, 2, 6)
    p_headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_hl = p_headline.add_run("AI PRODUCT MANAGER  |  BUSINESS TRANSFORMATION  |  ENTERPRISE AI STRATEGY")
    run_hl.font.size = Pt(10)
    run_hl.bold = True
    add_bottom_border(p_headline)

    # Helper for sections
    def add_section(title):
        p_title = doc.add_paragraph()
        format_spacing(p_title, 8, 4)
        run_title = p_title.add_run(title)
        run_title.font.size = Pt(10.5)
        run_title.bold = True
        add_bottom_border(p_title)
        return p_title

    # Summary
    add_section("SUMMARY")
    p_sum = doc.add_paragraph()
    format_spacing(p_sum, 0, 4)
    run_sum = p_sum.add_run(
        "AI Product Manager with 2+ years driving enterprise AI transformation at OSTTRA, a KKR portfolio company, and 4+ years total in enterprise automation and product delivery. "
        "Architect of AI systems processing 200,000+ monthly cases with >99% routing accuracy and designed prompts powering predictive sentiment analytics across 20,000+ client communications. "
        "Delivered 20+ FTE in annualized savings and 80%+ adoption across 1,500+ enterprise users through Product Strategy, Change Management, and rigorous adoption discipline. "
        "Expertise spans Vertex AI, RAG pipelines, Prompt Engineering, financial-services automation, and data-driven decision making. Open to AI Product Manager, AI Strategy, and Enterprise Transformation roles."
    )
    run_sum.font.size = Pt(9.5)

    # Work Experience
    add_section("WORK EXPERIENCE")

    # Job 1
    p_job1 = doc.add_paragraph()
    format_spacing(p_job1, 2, 1)
    p_job1.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
    run_j1 = p_job1.add_run("Associate, AI Strategy & Ops Excellence (DigiOps Programme)")
    run_j1.bold = True
    p_job1.add_run("\t")
    run_j1_date = p_job1.add_run("Apr 2025 – Present")
    run_j1_date.bold = True
    
    p_job1_meta = doc.add_paragraph()
    format_spacing(p_job1_meta, 0, 2)
    run_j1m = p_job1_meta.add_run("OSTTRA (KKR Portfolio Company) | Gurugram, India")
    run_j1m.italic = True

    bullets_job1 = [
        "Architected NEXUS global capacity intelligence platform consolidating 150+ spreadsheets using Python predictive modeling, Business Intelligence dashboards via Looker, and Data Analytics pipelines; drove 80%+ adoption across 1,500+ users and compressed planning cycles 45%",
        "Spearheaded Generative AI-powered Auto-Triaging ecosystem processing 200,000+ monthly cases with >99% routing accuracy; reclaimed 4.9 FTEs through Process Automation and scaled duplicate flagging to 15,206 monthly checks recovering 3.88 additional FTEs",
        "Led intelligent KYC and onboarding automation through API Integration with reference registry systems executing 751 monthly compliance screenings; shortened cycle time 55% from 18 to 8 days, reclaimed 1.0 FTE, and established 100% audit traceability",
        "Engineered autonomous failure management for 4,500+ monthly complex system alerts including LimitHub and TOC FIX disconnections; eliminated reactive triage burden by routing non-actionable items without human intervention",
        "Designed AI prompts and led change management for predictive sentiment analytics across 20,000+ client emails on MarkitWire and TriResolve; engineered dual-category classification (positive, negative, neutral) for Client Services and Product, enabling proactive client outreach before escalation accelerating resolution 40% and decreasing escalations 20%",
        "Contributed to the design and rollout of Dot, ODA, and Transformer Bot, RAG-based Conversational AI assistants powered by Vertex AI; supported the team that cut routine client information cases from 16% to 3%, freed 4+ FTE annually, and deployed across 24 knowledge initiatives"
    ]
    for b in bullets_job1:
        pb = doc.add_paragraph(style='List Bullet')
        format_spacing(pb, 0, 1)
        run_b = pb.add_run(b)
        run_b.font.size = Pt(9.5)

    # Job 2
    p_job2 = doc.add_paragraph()
    format_spacing(p_job2, 3, 1)
    p_job2.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
    run_j2 = p_job2.add_run("Senior Operational Excellence Analyst")
    run_j2.bold = True
    p_job2.add_run("\t")
    run_j2_date = p_job2.add_run("Feb 2024 – Dec 2024")
    run_j2_date.bold = True
    
    p_job2_meta = doc.add_paragraph()
    format_spacing(p_job2_meta, 0, 2)
    run_j2m = p_job2_meta.add_run("OSTTRA | Gurugram, India")
    run_j2m.italic = True

    bullets_job2 = [
        "Engineered SOX Compliance automation using Google Apps Script; eliminated 4.1 FTEs of manual audit work, achieved 100% compliance, and compressed audit cycle from 45 to 5 days",
        "Delivered 10+ continuous improvement initiatives applying Lean Six Sigma and Agile across client services and trading operations; unlocked 2+ FTE in quantitative savings",
        "Directed Microsoft to Google Workspace migration during Tel Aviv-to-India transition; removed VDI dependency, cut access latency 40%, and freed 0.4 FTEs"
    ]
    for b in bullets_job2:
        pb = doc.add_paragraph(style='List Bullet')
        format_spacing(pb, 0, 1)
        run_b = pb.add_run(b)
        run_b.font.size = Pt(9.5)

    # Job 3
    p_job3 = doc.add_paragraph()
    format_spacing(p_job3, 3, 1)
    p_job3.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
    run_j3 = p_job3.add_run("Credit Analyst Intern")
    run_j3.bold = True
    p_job3.add_run("\t")
    run_j3_date = p_job3.add_run("Feb 2023 – Apr 2023")
    run_j3_date.bold = True
    
    p_job3_meta = doc.add_paragraph()
    format_spacing(p_job3_meta, 0, 2)
    run_j3m = p_job3_meta.add_run("Kotak Mahindra Bank | Delhi NCR, India")
    run_j3m.italic = True

    pb3 = doc.add_paragraph(style='List Bullet')
    format_spacing(pb3, 0, 1)
    run_b3 = pb3.add_run("Evaluated 20+ MSME credit applications analyzing financial statements and working capital dynamics; delivered recommendations with 95%+ accuracy contributing to $3M+ underwriting decisions")
    run_b3.font.size = Pt(9.5)

    # Job 4
    p_job4 = doc.add_paragraph()
    format_spacing(p_job4, 3, 1)
    p_job4.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
    run_j4 = p_job4.add_run("React Native Developer")
    run_j4.bold = True
    p_job4.add_run("\t")
    run_j4_date = p_job4.add_run("Apr 2021 – May 2022")
    run_j4_date.bold = True
    
    p_job4_meta = doc.add_paragraph()
    format_spacing(p_job4_meta, 0, 2)
    run_j4m = p_job4_meta.add_run("RPQ IT Services | Noida (Sector 16), India")
    run_j4m.italic = True

    bullets_job4 = [
        "Developed 5+ production mobile applications using React Native with advanced state management and component architecture",
        "Managed full Software Development Life Cycle (SDLC) from requirement gathering through production deployment for iOS and Android platforms",
        "Architected cross-platform build pipelines and CI/CD integration enabling multi-platform releases"
    ]
    for b in bullets_job4:
        pb4 = doc.add_paragraph(style='List Bullet')
        format_spacing(pb4, 0, 1)
        run_b4 = pb4.add_run(b)
        run_b4.font.size = Pt(9.5)

    # Job 5
    p_job5 = doc.add_paragraph()
    format_spacing(p_job5, 3, 1)
    p_job5.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
    run_j5 = p_job5.add_run("React Native Developer")
    run_j5.bold = True
    p_job5.add_run("\t")
    run_j5_date = p_job5.add_run("Nov 2020 – Mar 2021")
    run_j5_date.bold = True
    
    p_job5_meta = doc.add_paragraph()
    format_spacing(p_job5_meta, 0, 2)
    run_j5m = p_job5_meta.add_run("Markup Designer | Noida (Sector 62), India")
    run_j5m.italic = True

    bullets_job5 = [
        "Engineered React Native mobile applications with REST API integration and asynchronous data handling",
        "Collaborated with design teams using Figma and Adobe XD translating wireframes into production-grade cross-platform code",
        "Optimized application performance achieving 40%+ faster load times through code splitting and lazy loading"
    ]
    for b in bullets_job5:
        pb5 = doc.add_paragraph(style='List Bullet')
        format_spacing(pb5, 0, 1)
        run_b5 = pb5.add_run(b)
        run_b5.font.size = Pt(9.5)

    # Job 6
    p_job6 = doc.add_paragraph()
    format_spacing(p_job6, 3, 1)
    p_job6.paragraph_format.tab_stops.add_tab_stop(Inches(7.0), docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)
    run_j6 = p_job6.add_run("Mobile Application Developer")
    run_j6.bold = True
    p_job6.add_run("\t")
    run_j6_date = p_job6.add_run("Jul 2020 – Nov 2020")
    run_j6_date.bold = True
    
    p_job6_meta = doc.add_paragraph()
    format_spacing(p_job6_meta, 0, 2)
    run_j6m = p_job6_meta.add_run("Sabpay | Remote, India (BCA Freelance Engagement)")
    run_j6m.italic = True

    bullets_job6 = [
        "Built mobile application user interfaces using React Native during freelance engagement while completing BCA",
        "Delivered UI/UX prototypes in Figma and Adobe XD shipping production builds on schedule",
        "Implemented OAuth authentication and payment gateway integration within the Sabpay ecosystem"
    ]
    for b in bullets_job6:
        pb6 = doc.add_paragraph(style='List Bullet')
        format_spacing(pb6, 0, 1)
        run_b6 = pb6.add_run(b)
        run_b6.font.size = Pt(9.5)

    # Education
    add_section("EDUCATION")
    edu_lines = [
        "Post Graduate Programme, Finance & Analytics — ICFAI Business School, Gurgaon  |  2022 – Apr 2024  |  CGPA: 9.0/10",
        "MBA, Finance & Analytics (Online, concurrent with OSTTRA role) — IFHE Hyderabad  |  2022 – Jul 2024  |  CGPA: 8.75/10",
        "Bachelor of Computer Applications — Birla Institute of Technology, Mesra  |  2017 – 2020  |  CGPA: 7.52/10"
    ]
    for el in edu_lines:
        pe = doc.add_paragraph(style='List Bullet')
        format_spacing(pe, 0, 1)
        parts = el.split('—')
        if len(parts) > 1:
            run_b = pe.add_run(parts[0])
            run_b.bold = True
            run_b.font.size = Pt(9.5)
            run_n = pe.add_run("—" + parts[1])
            run_n.font.size = Pt(9.5)
        else:
            run_n = pe.add_run(el)
            run_n.font.size = Pt(9.5)

    # Skills
    add_section("SKILLS")
    skills = [
        ("AI and LLM Stack: ", "Vertex AI, Google Gemini, Generative AI (GenAI), Machine Learning, RAG Pipelines, Vector Embeddings, Prompt Engineering, LLM Optimization, Conversational AI, Predictive Sentiment Analysis, Agentic Workflows, LLM-Powered Automation, ChatGPT"),
        ("Cloud and Data: ", "Python, SQL, Google BigQuery, ChromaDB, Google Cloud Platform (GCP), Looker Studio, Power BI, Data Analytics, Business Intelligence, Data Engineering, Metrics-Driven Analysis"),
        ("Product and Tools: ", "Salesforce (Administration, Dashboards, Case Management, Custom Fields), Jira, Confluence, Google AppSheet, Google Apps Script, REST APIs, API Integration"),
        ("Product Management: ", "Product Strategy, Requirements Definition, Roadmap Development, User Adoption, Feature Prioritization, Sprint Planning, KPI Tracking, Stakeholder Communication, Executive Reporting, Cross-Functional Leadership, End-to-End Ownership"),
        ("Methodologies: ", "Lean Six Sigma, Agile/Scrum, LUMA Design Thinking, Process Reengineering, Change Management, Stakeholder Enablement, Continuous Improvement, Data-Driven Decision Making, Cost Optimization"),
        ("Financial Domain: ", "OTC Markets, Reference Data, Trade Operations, Settlement, Reconciliation, KYC/AML, SOX Compliance, Regulatory Reporting, Audit Trail Management, Trade Failure Management, Financial Services Infrastructure, Credit Analysis")
    ]
    for lbl, val in skills:
        ps = doc.add_paragraph()
        format_spacing(ps, 0, 1)
        run_lbl = ps.add_run(lbl)
        run_lbl.bold = True
        run_lbl.font.size = Pt(9.5)
        run_val = ps.add_run(val)
        run_val.font.size = Pt(9.5)

    # Certifications
    add_section("CERTIFICATIONS")
    certs = [
        "Google Data Analytics Professional Certificate, Google (2022)",
        "Certified Data Scientist, Henry Harvin Education (2022)",
        "SQL Certification, DataCamp (2022)"
    ]
    for c in certs:
        pc = doc.add_paragraph(style='List Bullet')
        format_spacing(pc, 0, 1)
        run_c = pc.add_run(c)
        run_c.font.size = Pt(9.5)

    # Recognition
    add_section("RECOGNITION")
    recogs = [
        "Featured in Times of India, InspireTech 2025 — Enterprise AI and Automation Systems",
        "Highflyer Award, OSTTRA Q1 2026 — Looker analytics and Salesforce auto-triaging impact",
        "Trendsetter Award, OSTTRA Q1 2025 — SOX Compliance automation",
        "1st Prize, LUMA Design Thinking Challenge, Protiviti 2023",
        "Multiple Quarterly Spot Awards, OSTTRA 2024 – 2026"
    ]
    for r in recogs:
        pr = doc.add_paragraph(style='List Bullet')
        format_spacing(pr, 0, 1)
        run_r = pr.add_run(r)
        run_r.font.size = Pt(9.5)

    doc.save(output_path)
    print(f"Successfully generated DOCX to {output_path}")

ats_docx_path = os.path.join(output_dir, "damak-kumar-varshney-resume.docx")
build_docx(ats_docx_path)

# Generate TXT Version
ats_txt_path = os.path.join(output_dir, "damak-kumar-varshney-resume.txt")
print("Generating TXT version...")
txt_content = """DAMAK KUMAR VARSHNEY
Gurugram, India | +91-8077293920 | damakvarshney@gmail.com
linkedin.com/in/damak-varshney | damak-varshney.vercel.app

AI PRODUCT MANAGER | BUSINESS TRANSFORMATION | ENTERPRISE AI STRATEGY

=== SUMMARY ===
AI Product Manager with 2+ years driving enterprise AI transformation at OSTTRA, a KKR portfolio company, and 4+ years total in enterprise automation and product delivery. Architect of AI systems processing 200,000+ monthly cases with >99% routing accuracy and designed prompts powering predictive sentiment analytics across 20,000+ client communications. Delivered 20+ FTE in annualized savings and 80%+ adoption across 1,500+ enterprise users through Product Strategy, Change Management, and rigorous adoption discipline. Expertise spans Vertex AI, RAG pipelines, Prompt Engineering, financial-services automation, and data-driven decision making. Open to AI Product Manager, AI Strategy, and Enterprise Transformation roles.

=== WORK EXPERIENCE ===

Associate, AI Strategy & Ops Excellence (DigiOps Programme)
OSTTRA (KKR Portfolio Company) | Gurugram, India | Apr 2025 – Present
• Architected NEXUS global capacity intelligence platform consolidating 150+ spreadsheets using Python predictive modeling, Business Intelligence dashboards via Looker, and Data Analytics pipelines; drove 80%+ adoption across 1,500+ users and compressed planning cycles 45%.
• Spearheaded Generative AI-powered Auto-Triaging ecosystem processing 200,000+ monthly cases with >99% routing accuracy; reclaimed 4.9 FTEs through Process Automation and scaled duplicate flagging to 15,206 monthly checks recovering 3.88 additional FTEs.
• Led intelligent KYC and onboarding automation through API Integration with reference registry systems executing 751 monthly compliance screenings; shortened cycle time 55% from 18 to 8 days, reclaimed 1.0 FTE, and established 100% audit traceability.
• Engineered autonomous failure management for 4,500+ monthly complex system alerts including LimitHub and TOC FIX disconnections; eliminated reactive triage burden by routing non-actionable items without human intervention.
• Designed AI prompts and led change management for predictive sentiment analytics across 20,000+ client emails on MarkitWire and TriResolve; engineered dual-category classification (positive, negative, neutral) for Client Services and Product, enabling proactive client outreach before escalation accelerating resolution 40% and decreasing escalations 20%.
• Contributed to the design and rollout of Dot, ODA, and Transformer Bot, RAG-based Conversational AI assistants powered by Vertex AI; supported the team that cut routine client information cases from 16% to 3%, freed 4+ FTE annually, and deployed across 24 knowledge initiatives.

Senior Operational Excellence Analyst
OSTTRA | Gurugram, India | Feb 2024 – Dec 2024
• Engineered SOX Compliance automation using Google Apps Script; eliminated 4.1 FTEs of manual audit work, achieved 100% compliance, and compressed audit cycle from 45 to 5 days.
• Delivered 10+ continuous improvement initiatives applying Lean Six Sigma and Agile across client services and trading operations; unlocked 2+ FTE in quantitative savings.
• Directed Microsoft to Google Workspace migration during Tel Aviv-to-India transition; removed VDI dependency, cut access latency 40%, and freed 0.4 FTEs.

Credit Analyst Intern
Kotak Mahindra Bank | Delhi NCR, India | Feb 2023 – Apr 2023
• Evaluated 20+ MSME credit applications analyzing financial statements and working capital dynamics; delivered recommendations with 95%+ accuracy contributing to $3M+ underwriting decisions.

React Native Developer
RPQ IT Services | Noida (Sector 16), India | Apr 2021 – May 2022
• Developed 5+ production mobile applications using React Native with advanced state management and component architecture.
• Managed full Software Development Life Cycle (SDLC) from requirement gathering through production deployment for iOS and Android platforms.
• Architected cross-platform build pipelines and CI/CD integration enabling multi-platform releases.

React Native Developer
Markup Designer | Noida (Sector 62), India | Nov 2020 – Mar 2021
• Engineered React Native mobile applications with REST API integration and asynchronous data handling.
• Collaborated with design teams using Figma and Adobe XD translating wireframes into production-grade cross-platform code.
• Optimized application performance achieving 40%+ faster load times through code splitting and lazy loading.

Mobile Application Developer
Sabpay | Remote, India (BCA Freelance Engagement) | Jul 2020 – Nov 2020
• Built mobile application user interfaces using React Native during freelance engagement while completing BCA.
• Delivered UI/UX prototypes in Figma and Adobe XD shipping production builds on schedule.
• Implemented OAuth authentication and payment gateway integration within the Sabpay ecosystem.

=== EDUCATION ===
Post Graduate Programme, Finance & Analytics — ICFAI Business School, Gurgaon | 2022 – Apr 2024 | CGPA: 9.0/10
MBA, Finance & Analytics (Online, concurrent with OSTTRA role) — IFHE Hyderabad | 2022 – Jul 2024 | CGPA: 8.75/10
Bachelor of Computer Applications — Birla Institute of Technology, Mesra | 2017 – 2020 | CGPA: 7.52/10

=== SKILLS ===
• AI and LLM Stack: Vertex AI, Google Gemini, Generative AI (GenAI), Machine Learning, RAG Pipelines, Vector Embeddings, Prompt Engineering, LLM Optimization, Conversational AI, Predictive Sentiment Analysis, Agentic Workflows, LLM-Powered Automation, ChatGPT
• Cloud and Data: Python, SQL, Google BigQuery, ChromaDB, Google Cloud Platform (GCP), Looker Studio, Power BI, Data Analytics, Business Intelligence, Data Engineering, Metrics-Driven Analysis
• Product and Tools: Salesforce (Administration, Dashboards, Case Management, Custom Fields), Jira, Confluence, Google AppSheet, Google Apps Script, REST APIs, API Integration
• Product Management: Product Strategy, Requirements Definition, Roadmap Development, User Adoption, Feature Prioritization, Sprint Planning, KPI Tracking, Stakeholder Communication, Executive Reporting, Cross-Functional Leadership, End-to-End Ownership
• Methodologies: Lean Six Sigma, Agile/Scrum, LUMA Design Thinking, Process Reengineering, Change Management, Stakeholder Enablement, Continuous Improvement, Data-Driven Decision Making, Cost Optimization
• Financial Domain: OTC Markets, Reference Data, Trade Operations, Settlement, Reconciliation, KYC/AML, SOX Compliance, Regulatory Reporting, Audit Trail Management, Trade Failure Management, Financial Services Infrastructure, Credit Analysis

=== CERTIFICATIONS ===
• Google Data Analytics Professional Certificate, Google (2022)
• Certified Data Scientist, Henry Harvin Education (2022)
• SQL Certification, DataCamp (2022)

=== RECOGNITION ===
• Featured in Times of India, InspireTech 2025 — Enterprise AI and Automation Systems
• Highflyer Award, OSTTRA Q1 2026 — Looker analytics and Salesforce auto-triaging impact
• Trendsetter Award, OSTTRA Q1 2025 — SOX Compliance automation
• 1st Prize, LUMA Design Thinking Challenge, Protiviti 2023
• Multiple Quarterly Spot Awards, OSTTRA 2024 – 2026
"""

with open(ats_txt_path, "w", encoding="utf-8") as f:
    f.write(txt_content)
print(f"Successfully generated TXT to {ats_txt_path}")

